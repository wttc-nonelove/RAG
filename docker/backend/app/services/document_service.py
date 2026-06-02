"""
文档服务模块
功能：处理文档的上传、解析、向量化、知识图谱抽取和删除
支持格式：PDF、DOCX、TXT、MD（Markdown）、XLSX、XLS（Excel）、CSV
"""

import os
import pdfplumber
from docx import Document as DocxDocument
from sqlalchemy.ext.asyncio import AsyncSession

from app.dao import document_dao, file_store, chroma_repo
from app.models.document import Document
from app.utils.text_processor import clean_text, chunk_text
from app.utils.embedding import embedding_client
from app.utils.kg_extractor import extract as kg_extract
from app.dao import neo4j_repo


# 支持的文件类型映射
ALLOWED_TYPES = {"PDF": ".pdf", "DOCX": ".docx", "TXT": ".txt", "MD": ".md", "XLSX": ".xlsx", "XLS": ".xls", "CSV": ".csv"}


async def get_document_stats(db: AsyncSession) -> dict:
    """
    获取文档统计数据

    返回:
        dict: {
            "total": 总文件数,
            "type_counts": {"PDF": 5, "DOCX": 3, ...},
            "status_counts": {"completed": 10, "pending": 2, ...}
        }
    """
    from sqlalchemy import select, func
    # 总文件数
    total = (await db.execute(select(func.count()).select_from(Document))).scalar()
    # 各类型文件数
    type_result = await db.execute(
        select(Document.file_type, func.count()).group_by(Document.file_type)
    )
    type_counts = {row[0]: row[1] for row in type_result.all()}
    # 各状态文件数
    status_result = await db.execute(
        select(Document.parse_status, func.count()).group_by(Document.parse_status)
    )
    status_counts = {row[0]: row[1] for row in status_result.all()}
    return {
        "total": total,
        "type_counts": type_counts,
        "status_counts": status_counts,
    }


def _detect_type(filename: str) -> str:
    """根据文件扩展名检测文件类型"""
    ext = os.path.splitext(filename)[1].lower()
    for ft, e in ALLOWED_TYPES.items():
        if e == ext:
            return ft
    raise ValueError(f"不支持的文件类型: {ext}")


async def upload_and_parse(db: AsyncSession, file_content: bytes, filename: str, tag: str, user_id: int) -> Document:
    """
    上传文档并保存到数据库

    参数:
        db: 数据库会话
        file_content: 文件内容（字节）
        filename: 文件名
        tag: 标签（可选）
        user_id: 上传用户ID

    返回:
        Document: 创建的文档记录
    """
    file_type = _detect_type(filename)
    filepath = file_store.save_file(file_content, filename)
    doc = Document(
        filename=filename,
        file_type=file_type,
        file_size=len(file_content),
        file_path=filepath,
        tag=tag,
        parse_status="pending",
        uploaded_by=user_id,
    )
    doc = await document_dao.create(db, doc)
    await db.commit()
    return doc


async def parse_document(doc_id: int) -> None:
    """
    异步解析文档（后台任务）
    流程：提取文本 → 清洗 → 分块 → 向量化 → 入库 ChromaDB → 知识图谱抽取 → 入库 Neo4j

    参数:
        doc_id: 文档ID
    """
    from app.database import async_session
    async with async_session() as db:
        doc = await document_dao.get_by_id(db, doc_id)
        if not doc:
            return

        # 更新状态为"解析中"
        await document_dao.update_status(db, doc_id, "parsing")
        await db.commit()

        try:
            # 第一步：提取文本
            text = _extract_text(doc.file_path, doc.file_type)
            # 第二步：清洗文本
            text = clean_text(text)
            # 第三步：分块
            chunks = chunk_text(text)

            # 第四步：向量化并入库 ChromaDB
            embed_result = await embedding_client.encode_batch_from_db(db, chunks)
            embeddings = embed_result["embeddings"]
            embedding_tokens = embed_result["tokens_used"]
            ids = [f"doc{doc_id}_chunk{i}" for i in range(len(chunks))]
            metadatas = [{"doc_id": doc_id, "filename": doc.filename, "chunk_index": i} for i in range(len(chunks))]
            chroma_repo.add_batch(ids=ids, embeddings=embeddings, documents=chunks, metadatas=metadatas)

            # 记录 embedding token 消耗到独立表
            if embedding_tokens > 0:
                from app.dao import token_usage_dao
                # 获取 embedding 模型名称
                from app.dao import model_dao
                embed_config = await model_dao.get_default_config(db, "embedding")
                embed_model_name = embed_config.model_name if embed_config else "unknown"
                await token_usage_dao.create(
                    db, model_name=embed_model_name, model_type="embedding",
                    tokens_used=embedding_tokens, source_type="document",
                    source_id=doc_id, source_name=doc.filename
                )
                # 同时更新文档记录
                from sqlalchemy import update
                await db.execute(update(Document).where(Document.id == doc_id).values(embedding_tokens=embedding_tokens))
                await db.commit()

            # 第五步：知识图谱抽取
            try:
                # 从数据库读取提取参数配置
                from app.dao import config_dao
                config = await config_dao.get_all(db)
                config_dict = {c.config_key: c.config_value for c in config}
                kg_chunk_size = int(config_dict.get("kg_chunk_size", "3000"))
                kg_overlap = int(config_dict.get("kg_overlap", "500"))
                kg_min_chars = int(config_dict.get("kg_min_chars", "200"))

                # 分段提取知识图谱
                all_entities = []
                all_relations = []
                start = 0
                while start < len(text):
                    end = min(start + kg_chunk_size, len(text))
                    chunk = text[start:end]
                    if len(chunk) >= kg_min_chars:  # 至少指定字符数才提取
                        kg_data = await kg_extract(chunk, doc_id, doc.filename, db=db)
                        all_entities.extend(kg_data.get("entities", []))
                        all_relations.extend(kg_data.get("relations", []))
                    start += kg_chunk_size - kg_overlap
                    if end >= len(text):
                        break
                # 去重并存储
                seen_entities = set()
                for entity in all_entities:
                    key = entity["name"]
                    if key not in seen_entities:
                        seen_entities.add(key)
                        neo4j_repo.create_entity(entity["name"], entity.get("type", "概念"), entity.get("description", ""), doc_id)
                seen_relations = set()
                for rel in all_relations:
                    key = f"{rel['source']}|{rel['relation']}|{rel['target']}"
                    if key not in seen_relations:
                        seen_relations.add(key)
                        neo4j_repo.create_relation(rel["source"], rel["relation"], rel["target"])
            except Exception as e:
                import loguru
                loguru.logger.error(f"KG storage failed: {e}")

            # 更新状态为"已完成"
            await document_dao.update_status(db, doc_id, "completed")
        except Exception as e:
            # 更新状态为"失败"
            await document_dao.update_status(db, doc_id, "failed", str(e))

        await db.commit()


async def extract_kg_for_doc(doc_id: int) -> None:
    """
    手动触发指定文档的知识图谱抽取

    参数:
        doc_id: 文档ID
    """
    from app.database import async_session
    async with async_session() as db:
        doc = await document_dao.get_by_id(db, doc_id)
        if not doc:
            return
        try:
            text = _extract_text(doc.file_path, doc.file_type)
            text = clean_text(text)
            # 先删除旧的实体关系
            neo4j_repo.delete_by_doc_id(doc_id)
            # 重新抽取
            kg_data = await kg_extract(text[:3000], doc_id, doc.filename, db=db)
            for entity in kg_data.get("entities", []):
                neo4j_repo.create_entity(entity["name"], entity.get("type", "概念"), entity.get("description", ""), doc_id)
            for rel in kg_data.get("relations", []):
                neo4j_repo.create_relation(rel["source"], rel["relation"], rel["target"])
        except Exception as e:
            import loguru
            loguru.logger.error(f"KG re-extraction failed: {e}")


def _extract_text(filepath: str, file_type: str) -> str:
    """
    根据文件类型提取文本内容

    参数:
        filepath: 文件路径
        file_type: 文件类型（PDF/DOCX/TXT/MD/XLSX/XLS/CSV）

    返回:
        str: 提取的文本内容
    """
    if file_type == "PDF":
        with pdfplumber.open(filepath) as pdf:
            return "\n\n".join(page.extract_text() or "" for page in pdf.pages)
    elif file_type == "DOCX":
        doc = DocxDocument(filepath)
        return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
    elif file_type in ("TXT", "MD"):
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            return f.read()
    elif file_type in ("XLSX", "XLS"):
        from openpyxl import load_workbook
        wb = load_workbook(filepath, read_only=True, data_only=True)
        lines = []
        for sheet in wb.sheetnames:
            ws = wb[sheet]
            lines.append(f"【{sheet}】")
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join(str(cell) if cell is not None else "" for cell in row)
                if row_text.strip():
                    lines.append(row_text)
        wb.close()
        return "\n".join(lines)
    elif file_type == "CSV":
        import csv
        lines = []
        with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                row_text = "\t".join(row)
                if row_text.strip():
                    lines.append(row_text)
        return "\n".join(lines)
    raise ValueError(f"不支持的文件类型: {file_type}")


async def get_preview(db: AsyncSession, doc_id: int, page: int = 1, page_size: int = 50) -> dict:
    """
    获取文档预览（分页）

    参数:
        db: 数据库会话
        doc_id: 文档ID
        page: 页码
        page_size: 每页行数

    返回:
        dict: {filename, total_pages, current_page, page_size, content, has_next}
    """
    doc = await document_dao.get_by_id(db, doc_id)
    if not doc:
        raise ValueError("文档不存在")
    text = _extract_text(doc.file_path, doc.file_type)
    lines = text.split("\n")
    total_pages = (len(lines) + page_size - 1) // page_size
    start = (page - 1) * page_size
    end = start + page_size
    return {
        "filename": doc.filename,
        "total_pages": total_pages,
        "current_page": page,
        "page_size": page_size,
        "content": "\n".join(lines[start:end]),
        "has_next": end < len(lines),
    }


async def delete_document(db: AsyncSession, doc_id: int) -> None:
    """
    删除文档（级联清理：ChromaDB向量 + Neo4j图谱 + 本地文件 + 数据库记录）

    参数:
        db: 数据库会话
        doc_id: 文档ID
    """
    doc = await document_dao.get_by_id(db, doc_id)
    if not doc:
        return
    # 清理 ChromaDB 向量
    try:
        chroma_repo.delete_by_doc_id(doc_id)
    except Exception as e:
        import loguru
        loguru.logger.warning(f"ChromaDB delete failed for doc {doc_id}: {e}")
    # 清理 Neo4j 图谱
    try:
        neo4j_repo.delete_by_doc_id(doc_id)
    except Exception as e:
        import loguru
        loguru.logger.warning(f"Neo4j delete failed for doc {doc_id}: {e}")
    # 删除本地文件
    file_store.delete_file(doc.file_path)
    # 删除数据库记录
    await document_dao.delete_doc(db, doc_id)
    await db.commit()


async def cleanup_orphan_entities(db: AsyncSession) -> int:
    """
    清理没有对应文档的孤立实体

    返回:
        int: 删除的实体数量
    """
    # 获取所有文档 ID
    docs = await document_dao.get_all_ids(db)
    doc_ids = set(docs) if docs else set()
    # 清理 Neo4j 中的孤立实体
    deleted = neo4j_repo.cleanup_orphan_entities(list(doc_ids))
    return deleted
