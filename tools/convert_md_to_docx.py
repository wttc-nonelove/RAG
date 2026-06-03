from __future__ import annotations

import argparse
import re
import textwrap
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


BODY_FONT = "Microsoft YaHei"
CODE_FONT = "Consolas"
HEADING_BLUE = RGBColor(31, 77, 120)
HEADING_DARK = RGBColor(20, 47, 73)
TABLE_HEADER_FILL = "E8EEF5"
TABLE_BORDER = "B7C0CC"
AVAILABLE_WIDTH_CM = 16.5
DIAGRAM_DIR = Path("report_assets/generated_docx_diagrams")


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path("C:/Windows/Fonts/msyhbd.ttc" if bold else "C:/Windows/Fonts/msyh.ttc"),
        Path("C:/Windows/Fonts/simhei.ttf"),
        Path("C:/Windows/Fonts/simsun.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def text_size(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont) -> tuple[int, int]:
    box = draw.textbbox((0, 0), text, font=font)
    return box[2] - box[0], box[3] - box[1]


def wrap_label(text: str, width: int = 16) -> list[str]:
    text = text.replace("\\n", "\n")
    lines: list[str] = []
    for part in text.split("\n"):
        if len(part) <= width:
            lines.append(part)
        else:
            lines.extend(textwrap.wrap(part, width=width, break_long_words=True))
    return lines or [""]


def draw_centered_text(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    text: str,
    font: ImageFont.ImageFont,
    fill: str = "#1F2937",
    max_chars: int = 16,
) -> None:
    lines = wrap_label(text, max_chars)
    line_heights = [text_size(draw, line, font)[1] for line in lines]
    total_h = sum(line_heights) + max(0, len(lines) - 1) * 5
    y = box[1] + (box[3] - box[1] - total_h) // 2
    for line, h in zip(lines, line_heights):
        w, _ = text_size(draw, line, font)
        draw.text((box[0] + (box[2] - box[0] - w) // 2, y), line, font=font, fill=fill)
        y += h + 5


def draw_arrow(
    draw: ImageDraw.ImageDraw,
    start: tuple[int, int],
    end: tuple[int, int],
    color: str = "#4B5563",
    width: int = 3,
) -> None:
    draw.line((start, end), fill=color, width=width)
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 12
    p1 = (x2 - ux * size + px * size * 0.55, y2 - uy * size + py * size * 0.55)
    p2 = (x2 - ux * size - px * size * 0.55, y2 - uy * size - py * size * 0.55)
    draw.polygon([end, p1, p2], fill=color)


class DiagramRenderer:
    def __init__(self, base_dir: Path, doc_stem: str):
        self.base_dir = base_dir
        self.doc_stem = doc_stem
        self.count = 0
        self.base_dir.mkdir(parents=True, exist_ok=True)

    def render(self, lines: list[str]) -> Path:
        self.count += 1
        text = "\n".join(lines)
        out = self.base_dir / f"{self.doc_stem}_diagram_{self.count:02d}.png"
        first = next((line.strip() for line in lines if line.strip()), "")
        if first.startswith("erDiagram"):
            img = self._render_er(lines)
        elif first.startswith("sequenceDiagram"):
            img = self._render_sequence(lines)
        else:
            img = self._render_flowchart(lines)
        img.save(out)
        return out

    def _canvas(self, width: int, height: int, title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
        img = Image.new("RGB", (width, height), "#FFFFFF")
        draw = ImageDraw.Draw(img)
        title_font = load_font(28, True)
        draw.rounded_rectangle((22, 18, width - 22, height - 18), radius=18, outline="#CBD5E1", width=2, fill="#FFFFFF")
        draw.text((42, 34), title, font=title_font, fill="#0F172A")
        draw.line((42, 72, width - 42, 72), fill="#E2E8F0", width=2)
        return img, draw

    def _node_label(self, token: str, line: str) -> str:
        m = re.search(re.escape(token) + r'\s*(?:\{|\[)\s*"([^"]+)"\s*(?:\}|\])', line)
        if m:
            return m.group(1)
        return token

    def _draw_box(self, draw, xy, label, fill="#F8FAFC", outline="#2563EB", font_size=18, max_chars=14) -> None:
        draw.rounded_rectangle(xy, radius=14, fill=fill, outline=outline, width=2)
        draw_centered_text(draw, xy, label, load_font(font_size, True), max_chars=max_chars)

    def _render_flowchart(self, lines: list[str]) -> Image.Image:
        source = "\n".join(lines)
        title = "系统架构图" if "flowchart LR" in source else "业务流程图"
        edge_lines = [line.strip() for line in lines if "-->" in line]
        labels: dict[str, str] = {}
        edges: list[tuple[str, str, str]] = []

        for line in edge_lines:
            # Examples: A["x"] --> B["y"], H -- 否 --> I["x"]
            left, right = line.split("-->", 1)
            condition = ""
            cond_match = re.search(r"--\s*([^-\s][^-]*)\s*$", left)
            if cond_match:
                condition = cond_match.group(1).strip()
                left = left[: cond_match.start()]
            left_id = re.match(r"([A-Za-z0-9_]+)", left.strip())
            right_id = re.match(r"([A-Za-z0-9_]+)", right.strip())
            if not left_id or not right_id:
                continue
            a, b = left_id.group(1), right_id.group(1)
            left_label = self._node_label(a, left)
            right_label = self._node_label(b, right)
            if a not in labels or left_label != a:
                labels[a] = left_label
            if b not in labels or right_label != b:
                labels[b] = right_label
            edges.append((a, b, condition))

        if "flowchart LR" in source:
            width, height = 1600, 900
            img, draw = self._canvas(width, height, title)
            boxes = {
                "User": (60, 350, 250, 470),
                "Frontend": (335, 340, 575, 480),
                "Backend": (690, 340, 930, 480),
                "MySQL": (1090, 115, 1480, 205),
                "Chroma": (1090, 225, 1480, 315),
                "Neo4j": (1090, 335, 1480, 425),
                "Redis": (1090, 445, 1480, 535),
                "Files": (1090, 555, 1480, 645),
                "LLM": (1090, 665, 1480, 755),
            }
            for key, box in boxes.items():
                self._draw_box(draw, box, labels.get(key, key), fill="#EFF6FF" if key in {"Frontend", "Backend"} else "#F8FAFC")
            for a, b, _ in edges:
                if a in boxes and b in boxes:
                    start = (boxes[a][2], (boxes[a][1] + boxes[a][3]) // 2)
                    end = (boxes[b][0], (boxes[b][1] + boxes[b][3]) // 2)
                    draw_arrow(draw, start, end)
            return img

        order: list[str] = []
        incoming: dict[str, int] = {}
        for a, b, _ in edges:
            if a not in order:
                order.append(a)
            if b not in order:
                order.append(b)
            incoming[b] = incoming.get(b, 0) + 1
            incoming.setdefault(a, incoming.get(a, 0))

        levels = {node: 0 for node in order}
        for _ in range(max(1, len(order))):
            changed = False
            for a, b, _ in edges:
                new_level = levels.get(a, 0) + 1
                if new_level > levels.get(b, 0):
                    levels[b] = new_level
                    changed = True
            if not changed:
                break

        grouped: dict[int, list[str]] = {}
        for node in order:
            grouped.setdefault(levels.get(node, 0), []).append(node)

        width = 1500
        max_level = max(grouped) if grouped else 0
        height = max(780, 145 + (max_level + 1) * 105)
        img, draw = self._canvas(width, height, title)
        boxes: dict[str, tuple[int, int, int, int]] = {}
        y = 105
        for level in range(max_level + 1):
            row_nodes = grouped.get(level, [])
            count = len(row_nodes)
            if count == 1:
                row_width = 590
            elif count == 2:
                row_width = 430
            else:
                row_width = 340
            gap = 60
            total_width = count * row_width + max(0, count - 1) * gap
            x = (width - total_width) // 2
            for idx, node in enumerate(row_nodes):
                left = x + idx * (row_width + gap)
                box = (left, y, left + row_width, y + 68)
                is_decision = "?" in labels.get(node, "")
                if is_decision:
                    draw.rounded_rectangle(box, radius=28, fill="#FFF7ED", outline="#F97316", width=2)
                    draw_centered_text(draw, box, labels[node], load_font(17, True), max_chars=20)
                else:
                    self._draw_box(draw, box, labels.get(node, node), fill="#F8FAFC", font_size=17, max_chars=24)
                boxes[node] = box
            y += 105
        for a, b, condition in edges:
            if a not in boxes or b not in boxes:
                continue
            start = ((boxes[a][0] + boxes[a][2]) // 2, boxes[a][3])
            end = ((boxes[b][0] + boxes[b][2]) // 2, boxes[b][1])
            draw_arrow(draw, start, end)
            if condition:
                mid = ((start[0] + end[0]) // 2 + 16, (start[1] + end[1]) // 2 - 18)
                draw.rounded_rectangle((mid[0] - 4, mid[1] - 2, mid[0] + 32, mid[1] + 22), radius=4, fill="#FFFFFF")
                draw.text(mid, condition, font=load_font(15, True), fill="#EA580C")
        return img

    def _render_sequence(self, lines: list[str]) -> Image.Image:
        participants: list[tuple[str, str]] = []
        messages: list[tuple[str, str, str, bool]] = []
        for line in lines:
            stripped = line.strip()
            m = re.match(r"participant\s+(\w+)\s+as\s+(.+)", stripped)
            if m:
                participants.append((m.group(1), m.group(2)))
                continue
            msg = re.match(r"(\w+)(-->>|->>)(\w+):\s*(.+)", stripped)
            if msg:
                messages.append((msg.group(1), msg.group(3), msg.group(4), msg.group(2).startswith("--")))

        width = max(1300, 150 + len(participants) * 170)
        height = max(840, 190 + len(messages) * 58)
        img, draw = self._canvas(width, height, "交互时序图")
        x_positions = {}
        step = (width - 140) // max(len(participants) - 1, 1)
        top = 105
        bottom = height - 60
        for idx, (pid, label) in enumerate(participants):
            x = 70 + idx * step
            x_positions[pid] = x
            box = (x - 68, top, x + 68, top + 54)
            self._draw_box(draw, box, label, fill="#EFF6FF", font_size=15, max_chars=10)
            draw.line((x, top + 54, x, bottom), fill="#CBD5E1", width=2)

        y = top + 90
        for src, dst, label, dashed in messages:
            if src not in x_positions or dst not in x_positions:
                continue
            x1, x2 = x_positions[src], x_positions[dst]
            if src == dst:
                draw.rounded_rectangle((x1 + 10, y - 12, x1 + 115, y + 24), radius=8, outline="#64748B", width=2)
                draw.text((x1 + 122, y - 10), label, font=load_font(13), fill="#334155")
            else:
                if dashed:
                    for seg in range(abs(x2 - x1) // 18):
                        sx = min(x1, x2) + seg * 18
                        draw.line((sx, y, sx + 9, y), fill="#64748B", width=2)
                    draw_arrow(draw, (x2 - 18 if x2 > x1 else x2 + 18, y), (x2, y), "#64748B", 2)
                else:
                    draw_arrow(draw, (x1, y), (x2, y), "#334155", 2)
                tx = min(x1, x2) + 8
                draw.text((tx, y - 22), label, font=load_font(13), fill="#334155")
            y += 58
        return img

    def _render_er(self, lines: list[str]) -> Image.Image:
        entities: dict[str, list[str]] = {}
        relationships: list[tuple[str, str, str]] = []
        current = None
        for raw in lines:
            line = raw.strip()
            rel = re.match(r"(\w+)\s+\S+\s+(\w+)\s+:\s+(.+)", line)
            if rel:
                relationships.append((rel.group(1), rel.group(2), rel.group(3)))
                continue
            start = re.match(r"(\w+)\s+\{", line)
            if start:
                current = start.group(1)
                entities[current] = []
                continue
            if line == "}":
                current = None
                continue
            if current and line:
                entities[current].append(line)

        width, height = 1650, 1260
        img, draw = self._canvas(width, height, "数据库实体关系图")
        positions = {
            "users": (70, 125, 400, 345),
            "documents": (510, 105, 890, 435),
            "conversations": (1030, 125, 1360, 325),
            "messages": (1030, 425, 1360, 720),
            "model_providers": (70, 610, 450, 825),
            "model_configs": (530, 590, 910, 865),
            "model_presets": (1030, 830, 1410, 1135),
        }
        for name, box in positions.items():
            fields = entities.get(name, [])
            draw.rounded_rectangle(box, radius=12, fill="#F8FAFC", outline="#2563EB", width=2)
            header = (box[0], box[1], box[2], box[1] + 38)
            draw.rounded_rectangle(header, radius=12, fill="#DBEAFE", outline="#2563EB", width=0)
            draw.text((box[0] + 14, box[1] + 8), name, font=load_font(17, True), fill="#1E3A8A")
            y = box[1] + 48
            for field in fields[:11]:
                draw.text((box[0] + 14, y), field, font=load_font(13), fill="#334155")
                y += 22
            if len(fields) > 11:
                draw.text((box[0] + 14, y), f"... 另 {len(fields) - 11} 项", font=load_font(13), fill="#64748B")

        for a, b, label in relationships:
            if a not in positions or b not in positions:
                continue
            ba, bb = positions[a], positions[b]
            start = (ba[2], (ba[1] + ba[3]) // 2)
            end = (bb[0], (bb[1] + bb[3]) // 2)
            if start[0] > end[0]:
                start = (ba[0], (ba[1] + ba[3]) // 2)
                end = (bb[2], (bb[1] + bb[3]) // 2)
            draw_arrow(draw, start, end, "#64748B", 2)
            mid = ((start[0] + end[0]) // 2, (start[1] + end[1]) // 2)
            draw.rectangle((mid[0] - 45, mid[1] - 15, mid[0] + 45, mid[1] + 12), fill="#FFFFFF")
            draw.text((mid[0] - 38, mid[1] - 13), label, font=load_font(12), fill="#475569")

        if relationships:
            summary_x, summary_y = 70, 930
            draw.text((summary_x, summary_y), "主要关系", font=load_font(17, True), fill="#0F172A")
            y = summary_y + 36
            for a, b, label in relationships:
                text = f"{a} -> {b} : {label}"
                draw.rounded_rectangle((summary_x, y, summary_x + 780, y + 32), radius=8, fill="#F1F5F9", outline="#CBD5E1")
                draw.text((summary_x + 12, y + 6), text, font=load_font(13), fill="#334155")
                y += 40
        return img


def set_run_font(run, name: str = BODY_FONT, size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_style_font(style, name: str, size: float, color: RGBColor | None = None, bold: bool | None = None) -> None:
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    style.font.size = Pt(size)
    if color:
        style.font.color.rgb = color
    if bold is not None:
        style.font.bold = bold


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = TABLE_BORDER) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_fixed_width(table, column_widths_cm: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Cm(column_widths_cm[idx])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(column_widths_cm[idx] / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.3)
    section.right_margin = Cm(2.2)
    section.header_distance = Cm(1.25)
    section.footer_distance = Cm(1.25)

    styles = doc.styles
    set_style_font(styles["Normal"], BODY_FONT, 10.5)
    styles["Normal"].paragraph_format.line_spacing = 1.2
    styles["Normal"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Heading 1"], BODY_FONT, 18, HEADING_DARK, True)
    styles["Heading 1"].paragraph_format.space_before = Pt(16)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)

    set_style_font(styles["Heading 2"], BODY_FONT, 15, HEADING_BLUE, True)
    styles["Heading 2"].paragraph_format.space_before = Pt(14)
    styles["Heading 2"].paragraph_format.space_after = Pt(6)

    set_style_font(styles["Heading 3"], BODY_FONT, 12.5, HEADING_BLUE, True)
    styles["Heading 3"].paragraph_format.space_before = Pt(10)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)

    for style_name in ("Heading 4", "Heading 5", "Heading 6"):
        set_style_font(styles[style_name], BODY_FONT, 11.5, HEADING_BLUE, True)
        styles[style_name].paragraph_format.space_before = Pt(8)
        styles[style_name].paragraph_format.space_after = Pt(3)

    if "List Bullet" in styles:
        set_style_font(styles["List Bullet"], BODY_FONT, 10.5)
        styles["List Bullet"].paragraph_format.space_after = Pt(3)
    if "List Number" in styles:
        set_style_font(styles["List Number"], BODY_FONT, 10.5)
        styles["List Number"].paragraph_format.space_after = Pt(3)


def split_inline_code(text: str) -> list[tuple[str, bool]]:
    parts: list[tuple[str, bool]] = []
    pos = 0
    for match in re.finditer(r"`([^`]+)`", text):
        if match.start() > pos:
            parts.append((text[pos:match.start()], False))
        parts.append((match.group(1), True))
        pos = match.end()
    if pos < len(text):
        parts.append((text[pos:], False))
    return parts or [("", False)]


def clean_inline_markup(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    return text


def add_markdown_paragraph(doc: Document, text: str, style: str | None = None) -> None:
    para = doc.add_paragraph(style=style)
    for part, is_code in split_inline_code(clean_inline_markup(text)):
        run = para.add_run(part)
        if is_code:
            set_run_font(run, CODE_FONT, 9.5)
            run.font.color.rgb = RGBColor(95, 99, 104)
        else:
            set_run_font(run, BODY_FONT, 10.5)


def is_table_separator(line: str) -> bool:
    return bool(re.match(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$", line))


def parse_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [clean_inline_markup(cell.strip()) for cell in stripped.split("|")]


def compute_column_widths(rows: list[list[str]]) -> list[float]:
    column_count = max(len(r) for r in rows)
    weights = []
    for idx in range(column_count):
        max_len = max((len(r[idx]) if idx < len(r) else 0) for r in rows)
        weights.append(max(1.0, min(float(max_len), 32.0)))
    total = sum(weights)
    widths = [max(1.6, AVAILABLE_WIDTH_CM * w / total) for w in weights]
    scale = AVAILABLE_WIDTH_CM / sum(widths)
    return [round(w * scale, 2) for w in widths]


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    column_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=column_count)
    table.style = "Table Grid"
    set_table_borders(table)
    widths = compute_column_widths(rows)
    set_table_fixed_width(table, widths)

    for row_idx, row in enumerate(rows):
        for col_idx in range(column_count):
            cell = table.cell(row_idx, col_idx)
            text = row[col_idx] if col_idx < len(row) else ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, TABLE_HEADER_FILL)
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
            run = para.add_run(text)
            set_run_font(run, BODY_FONT, 9.2 if column_count >= 5 else 9.8, bold=(row_idx == 0))
    doc.add_paragraph()


def add_code_block(doc: Document, code_lines: list[str]) -> None:
    text = "\n".join(code_lines).rstrip()
    if not text:
        return
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.45)
    para.paragraph_format.right_indent = Cm(0.15)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(8)
    para.paragraph_format.line_spacing = 1.0
    run = para.add_run(text)
    set_run_font(run, CODE_FONT, 8.5)
    run.font.color.rgb = RGBColor(68, 68, 68)


def add_diagram(doc: Document, image_path: Path) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(16.2))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)
    cap_run = caption.add_run("图示")
    set_run_font(cap_run, BODY_FONT, 9)
    cap_run.font.color.rgb = RGBColor(100, 116, 139)


def add_markdown_image(doc: Document, alt_text: str, image_path: Path) -> None:
    if not image_path.exists():
        add_markdown_paragraph(doc, f"[图片缺失：{alt_text} - {image_path}]")
        return
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(16.2))
    if alt_text:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.paragraph_format.space_after = Pt(8)
        cap_run = caption.add_run(alt_text)
        set_run_font(cap_run, BODY_FONT, 9)
        cap_run.font.color.rgb = RGBColor(100, 116, 139)


def add_footer(doc: Document, label: str) -> None:
    footer = doc.sections[0].footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(label)
    set_run_font(run, BODY_FONT, 8.5)
    run.font.color.rgb = RGBColor(102, 102, 102)


def convert(md_path: Path, docx_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    configure_document(doc)
    add_footer(doc, md_path.stem)
    diagram_renderer = DiagramRenderer(DIAGRAM_DIR, md_path.stem)

    in_code = False
    code_lang = ""
    code_lines: list[str] = []
    paragraph_buffer: list[str] = []
    i = 0

    def flush_paragraph() -> None:
        nonlocal paragraph_buffer
        if paragraph_buffer:
            add_markdown_paragraph(doc, " ".join(p.strip() for p in paragraph_buffer if p.strip()))
            paragraph_buffer = []

    while i < len(lines):
        line = lines[i]

        if line.startswith("```"):
            flush_paragraph()
            if in_code:
                if code_lang == "mermaid":
                    image_path = diagram_renderer.render(code_lines)
                    add_diagram(doc, image_path)
                else:
                    add_code_block(doc, code_lines)
                code_lines = []
                code_lang = ""
                in_code = False
            else:
                in_code = True
                code_lang = line.strip().lstrip("`").strip().lower()
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            flush_paragraph()
            i += 1
            continue

        image = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if image:
            flush_paragraph()
            alt_text = image.group(1).strip()
            raw_path = image.group(2).strip()
            image_path = Path(raw_path)
            if not image_path.is_absolute():
                image_path = md_path.parent / image_path
            add_markdown_image(doc, alt_text, image_path)
            i += 1
            continue

        if "|" in line and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_paragraph()
            table_rows = [parse_table_row(line)]
            i += 2
            while i < len(lines) and "|" in lines[i] and lines[i].strip():
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        if heading:
            flush_paragraph()
            level = len(heading.group(1))
            text = clean_inline_markup(heading.group(2).strip())
            doc.add_heading(text, min(level, 4))
            i += 1
            continue

        bullet = re.match(r"^\s*[-*]\s+(.+)$", line)
        if bullet:
            flush_paragraph()
            add_markdown_paragraph(doc, bullet.group(1).strip(), style="List Bullet")
            i += 1
            continue

        ordered = re.match(r"^\s*\d+\.\s+(.+)$", line)
        if ordered:
            flush_paragraph()
            add_markdown_paragraph(doc, ordered.group(1).strip(), style="List Number")
            i += 1
            continue

        paragraph_buffer.append(line)
        i += 1

    flush_paragraph()
    if code_lines:
        add_code_block(doc, code_lines)

    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("inputs", nargs="+", type=Path)
    args = parser.parse_args()

    for input_path in args.inputs:
        output_path = input_path.with_suffix(".docx")
        convert(input_path, output_path)
        print(output_path)


if __name__ == "__main__":
    main()
